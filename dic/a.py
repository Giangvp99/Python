import sys
import re
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor

words = []
with open("words.txt") as f:
    for line in f:
        words.append(line.strip())
for w in words:

    headers = requests.utils.default_headers()
    headers.update(
        {
            "User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.114 Safari/537.36",
        }
    )
    response = requests.get(
        "https://www.oxfordlearnersdictionaries.com/definition/english/" + w,
        headers=headers,
    )

    soup = BeautifulSoup(response.content, "html.parser")
    num = 1
    try:
        doc = Document()

        word = soup.find("h1", class_="headword").get_text()
        phon = soup.find("span", class_="phon").get_text()
        doc.add_heading(word + "    " + phon)

        suggest = "(" + soup.select(".top-container .pos")[0].get_text() + ")   "

        senses_ = soup.select("div.entry > ol")[0]
        for sh in senses_.select("span.shcut-g"):
            # shcut
            shcut = sh.h2
            font_shcut = (
                doc.add_paragraph().add_run("\n" + shcut.get_text().capitalize()).font
            )
            font_shcut.color.rgb = RGBColor(51, 153, 255)
            font_shcut.size = Pt(12)

            for sense in sh.select(".shcut-g>.sense"):
                grammar = (
                    sense.find("span", hclass="grammar").get_text()
                    if sense.select(".sense .grammar")
                    else ""
                )
                use = (
                    sense.select(".sense>.use")[0].get_text() + " "
                    if sense.select(".sense>.use")
                    else ""
                )
                cf = (
                    sense.select(".sense>.cf")[0].get_text() + " "
                    if sense.select(".sense>.cf")
                    else ""
                )
                labels = (
                    sense.select(".sense>.labels")[0].get_text() + " "
                    if sense.select(".sense>.labels")
                    else ""
                )
                _def = sense.select(".def")[0]
                # add to suggest
                suggest += _def.get_text() + ";     "
                # def
                para_def = doc.add_paragraph()
                run_def_n = para_def.add_run(str(num) + ".    ")
                run_def_n.bold = True
                run_def_n.size = Pt(12.5)
                num += 1

                # run_def = para_def.add_run(use + grammar + cf + labels)
                # run_def.italic = True
                # run_def.font.color.rgb = RGBColor(135, 135, 135)
                para_def.add_run(grammar).font.color.rgb = RGBColor(135, 135, 135)
                para_def.add_run(cf).bold = True
                run_def_l = para_def.add_run(labels)
                run_def_l.italic = True
                run_def_l.font.color.rgb = RGBColor(135, 135, 135)
                para_def.add_run(use)
                para_def.add_run(
                    ("\t" if use or grammar or cf or labels else "") + _def.get_text()
                ).italic = True
                # synonym /
                xrefs = sense.select(".xrefs")[0] if sense.select(".xrefs") else []
                prefix = xrefs.select(".prefix")[0] if xrefs else []
                if xrefs and prefix.get_text().upper() == "SYNONYM":
                    para = doc.add_paragraph(prefix.get_text().upper() + " ")
                    for r in xrefs.select(".Ref .xh"):
                        font1 = para.add_run(r.get_text()).font
                        font1.color.rgb = RGBColor(0, 128, 255)
                elif xrefs and prefix.get_text().upper() == "OPPOSITE":
                    para = doc.add_paragraph(prefix.get_text().upper() + " ")
                    for r in xrefs.select(".Ref .xh"):
                        font1 = para.add_run(r.get_text()).font
                        font1.color.rgb = RGBColor(0, 128, 255)
                # examples
                for e in sense.select(".sense>.examples>li"):
                    e_cf = e.select(".cf")[0].get_text() if e.select(".cf") else ""
                    e_x = e.select(".x")[0] if e.select(".x") else ""
                    para_e = doc.add_paragraph(style="List Bullet 2")
                    para_e.add_run(e_cf + ("  " if e_cf else "")).bold = True
                    for content in e_x.contents:
                        if re.search('<.*?("cl")+>', str(content)):
                            abc = re.sub("<.*?>", "", str(content))
                            run_e = para_e.add_run(abc)
                            run_e.bold = True
                        else:
                            abc = re.sub("<.*?>", "", str(content))
                            para_e.add_run(abc)

                # collocations
                c = [
                    col
                    for col in sense.select(".collapse")
                    if col.select(".box_title")[0].get_text()
                    == "Oxford Collocations Dictionary"
                ]
                if c:
                    body = c[0].select("span.body")[0]
                    for u in body.select("span.unbox"):
                        doc.add_paragraph(style="List 2").add_run(
                            u.get_text()
                        ).bold = True
                        for li in u.next_sibling.find_all("li", class_="li"):
                            if li.get_text() != "…":
                                doc.add_paragraph(li.get_text(), style="List Bullet 3")
        else:
            for sense in senses_.select("ol>.sense"):
                grammar = (
                    sense.find("span", hclass="grammar").get_text()
                    if sense.select(".sense .grammar")
                    else ""
                )
                use = (
                    sense.select(".sense>.use")[0].get_text() + " "
                    if sense.select(".sense>.use")
                    else ""
                )
                cf = (
                    sense.select(".sense>.sensetop>.cf")[0].get_text() + " "
                    if sense.select(".sense>.sensetop>.cf")
                    else ""
                )
                labels = (
                    sense.select(".sense>.labels")[0].get_text() + " "
                    if sense.select(".sense>.labels")
                    else ""
                )
                _def = sense.select(".def")[0]
                # add to suggest
                suggest += _def.get_text() + ";     "
                # def
                para_def = doc.add_paragraph()
                run_def_n = para_def.add_run(str(num) + ".    ")
                run_def_n.bold = True
                run_def_n.size = Pt(12.5)
                num += 1

                # run_def = para_def.add_run(use + grammar + cf + labels)
                # run_def.italic = True
                # run_def.font.color.rgb = RGBColor(135, 135, 135)
                para_def.add_run(grammar).font.color.rgb = RGBColor(135, 135, 135)
                para_def.add_run(cf).bold = True
                run_def_l = para_def.add_run(labels)
                run_def_l.italic = True
                run_def_l.font.color.rgb = RGBColor(135, 135, 135)
                para_def.add_run(use)
                para_def.add_run(
                    ("    " if use or grammar or cf or labels else "") + _def.get_text()
                ).italic = True
                # synonym /
                xrefs = sense.select(".xrefs")[0] if sense.select(".xrefs") else []
                prefix = xrefs.select(".prefix")[0] if xrefs else []
                if xrefs and prefix.get_text().upper() == "SYNONYM":
                    para = doc.add_paragraph(prefix.get_text().upper() + " ")
                    for r in xrefs.select(".Ref .xh"):
                        font1 = para.add_run(r.get_text()).font
                        font1.color.rgb = RGBColor(0, 128, 255)
                elif xrefs and prefix.get_text().upper() == "OPPOSITE":
                    para = doc.add_paragraph(prefix.get_text().upper() + " ")
                    for r in xrefs.select(".Ref .xh"):
                        font1 = para.add_run(r.get_text()).font
                        font1.color.rgb = RGBColor(0, 128, 255)

                # examples
                for e in sense.select(".sense>.examples>li"):
                    e_cf = e.select(".cf")[0].get_text() if e.select(".cf") else ""
                    e_x = e.select(".x")[0] if e.select(".x") else ""
                    para_e = doc.add_paragraph(style="List Bullet 2")
                    para_e.add_run(e_cf + ("  " if e_cf else "")).bold = True
                    for content in e_x.contents:
                        if re.search('<.*?("cl")+>', str(content)):
                            abc = re.sub("<.*?>", "", str(content))
                            run_e = para_e.add_run(abc)
                            run_e.bold = True
                        else:
                            abc = re.sub("<.*?>", "", str(content))
                            para_e.add_run(abc)
                # collocations
                c = [
                    col
                    for col in sense.select(".collapse")
                    if col.select(".box_title")[0].get_text()
                    == "Oxford Collocations Dictionary"
                ]
                if c:
                    body = c[0].select("span.body")[0]
                    for u in body.select("span.unbox"):
                        doc.add_paragraph(style="List 2").add_run(
                            u.get_text()
                        ).bold = True
                        for li in u.next_sibling.find_all("li", class_="li"):
                            if li.get_text() != "…":
                                doc.add_paragraph(li.get_text(), style="List Bullet 3")
                    else:
                        for p in body.select("span.p"):
                            run_p = doc.add_paragraph(style="List 2").add_run(
                                p.select("span.eb")[0].next_sibling.string
                            )
                            run_p.italic = True
                            run_p.bold = True
                            for li in p.select("ul>li"):
                                if li.get_text() != "…":
                                    doc.add_paragraph(
                                        li.get_text(), style="List Bullet 3"
                                    )

        if soup.select(".idioms"):
            font = doc.add_paragraph().add_run("Idioms").font
            font.bold = True
            font.size = Pt(14)
            font.color.rgb = RGBColor(37, 223, 217)
            for idiom in soup.select(".idioms > .idm-g"):
                idm = idiom.find("span", class_="idm")
                doc.add_paragraph().add_run(idm.get_text()).bold = True
                for sense in idiom.select("ol>li.sense"):
                    _def = sense.find("span", class_="def")
                    labels = (
                        sense.select(".sense .labels")[0].get_text() + " "
                        if sense.select(".sense .labels")
                        else ""
                    )
                    doc.add_paragraph(labels + _def.get_text(), style="List 2")
                    for exam in sense.find_all("span", class_="x"):
                        doc.add_paragraph(
                            exam.get_text().strip(), style="List Bullet 3"
                        )
                    xrefs = sense.select(".xrefs")[0] if sense.select(".xrefs") else []
                    prefix = xrefs.select(".prefix")[0] if xrefs else []
                    if xrefs and prefix.get_text().upper() == "SYNONYM":
                        para = doc.add_paragraph(prefix.get_text().upper() + " ")
                        for r in xrefs.select(".Ref .xh"):
                            font1 = para.add_run(r.get_text()).font
                            font1.color.rgb = RGBColor(0, 128, 255)
        if soup.select(".phrasal_verb_links"):
            font = doc.add_paragraph().add_run("Phrasal Verbs").font
            font.bold = True
            font.size = Pt(14)
            font.color.rgb = RGBColor(37, 223, 217)
            for li in soup.select(".phrasal_verb_links .pvrefs>li"):
                xh = li.select(".xh")[0]
                doc.add_paragraph(xh.get_text(), style="List Bullet 2")
        doc.add_page_break()
        doc.add_paragraph(suggest)
        doc.save(w + ".docx")
    except Exception as e:
        print("e: " + w)
        print(e)
