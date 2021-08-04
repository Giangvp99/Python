from docx import Document
from docx.shared import Inches
from docx.shared import Pt

document = Document()
paragraph = document.add_paragraph(
    "Indentation is the horizontal space between a paragraph and edge of its container"
)
paragraph_format = paragraph.paragraph_format
paragraph_format.first_line_indent = Inches(0.5)
run = document.add_paragraph().add_run(
    "run exists in the style inheritance hierarchy and by default inherits its character formatting from that hierarchy"
)
font = run.font
font.size = Pt(16)
font.italic = True
document.save("test.docx")
